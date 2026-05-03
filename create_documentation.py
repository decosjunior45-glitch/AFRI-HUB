#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Script pour générer la documentation complète du projet AFRI-HUB en format .docx
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

# Créer un document Word
doc = Document()

# Configuration des styles
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ============== TITRE PRINCIPAL ==============
title = doc.add_heading('AFRI-HUB', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_format = title.runs[0]
title_format.font.size = Pt(28)
title_format.font.bold = True
title_format.font.color.rgb = RGBColor(5, 150, 105)

subtitle = doc.add_paragraph('Plateforme Multi-Pays Africaine - Documentation Complète')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_format = subtitle.runs[0]
subtitle_format.font.size = Pt(14)
subtitle_format.font.italic = True
subtitle_format.font.color.rgb = RGBColor(107, 114, 128)

doc.add_paragraph()

# ============== TABLE OF CONTENTS ==============
doc.add_heading('📋 Table des matières', level=1)
toc_items = [
    '1. Vue d\'ensemble du projet',
    '2. Architecture globale',
    '3. Structure des fichiers',
    '4. Fichiers backend',
    '5. Fichiers frontend',
    '6. Fichiers de configuration',
    '7. Explications détaillées du code'
]
for item in toc_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ============== 1. VUE D'ENSEMBLE ==============
doc.add_heading('1. Vue d\'ensemble du projet', level=1)

doc.add_heading('Qu\'est-ce que AFRI-HUB ?', level=2)
doc.add_paragraph(
    'AFRI-HUB est une plateforme multi-pays africaine qui permet aux utilisateurs de:\n'
    '• Se connecter via authentification JWT\n'
    '• Accéder à leur espace via un sous-domaine dédié par pays (ex: senegal.localhost)\n'
    '• Gérer des items (tâches) isolés par pays\n'
    '• Voir leur pays, code téléphone et drapeau emoji en interface'
)

doc.add_heading('Objectif principal', level=2)
objectives = [
    'Créer une plateforme 10 pays africains (au lieu de 2 initialement)',
    'Interface minimaliste et moderne de sélection de pays',
    'Affichage des drapeaux emoji colorés',
    'Codes téléphone pour chaque pays',
    'Architecture multi-tenant via subdomains',
    'Interface SaaS professionnelle avec Tailwind CSS'
]
for obj in objectives:
    doc.add_paragraph(obj, style='List Bullet')

doc.add_page_break()

# ============== 2. ARCHITECTURE ==============
doc.add_heading('2. Architecture globale', level=1)

doc.add_paragraph('AFRI-HUB utilise une architecture client-serveur avec MongoDB:')

# Schéma
arch_table = doc.add_table(rows=4, cols=2)
arch_table.style = 'Light Grid Accent 1'
arch_cells = arch_table.rows[0].cells
arch_cells[0].text = 'Composant'
arch_cells[1].text = 'Détails'

arch_data = [
    ('Frontend (React)', 'Port 5173, URL racine: afri-hub.localhost:5173 (sélection), URL subdomain: {country}.localhost:5173 (app)'),
    ('Backend (Express)', 'Port 4000, Endpoints API /api/countries, /api/items, /api/auth, Détection pays via subdomain'),
    ('Base de données (MongoDB)', 'Collections: countries (10 pays), links, users, items')
]

for item, detail in arch_data:
    row_cells = arch_table.add_row().cells
    row_cells[0].text = item
    row_cells[1].text = detail

doc.add_heading('Flux utilisateur', level=2)
flux = [
    '1. Utilisateur → afri-hub.localhost:5173 (Page de sélection)',
    '2. Tape un pays → "Sénégal" trouvé',
    '3. Clique → Redirection à senegal.localhost:5173',
    '4. Backend détecte "senegal" du subdomain',
    '5. Retourne données du Sénégal',
    '6. Login/Register avec JWT',
    '7. Items isolés par pays + utilisateur'
]
for step in flux:
    doc.add_paragraph(step, style='List Number')

doc.add_page_break()

# ============== 3. STRUCTURE DES FICHIERS ==============
doc.add_heading('3. Structure des fichiers', level=1)

structure_text = """
AFRI-HUB/
├── src/                          # Backend TypeScript
│   ├── server.ts                # Point d'entrée principal (Express)
│   ├── types/
│   │   └── country.ts           # Interfaces TypeScript (Country, Link)
│   ├── utils/
│   │   ├── db.ts                # Connexion MongoDB
│   │   ├── subdomain.ts         # Extraction code pays du host header
│   │   ├── seedData.ts          # Initialisation 10 pays africains
│   │   └── errorHandler.ts      # Gestion globale erreurs
│   └── controllers/
│       ├── countriesController.ts # Endpoints /api/countries
│       ├── itemsController.ts     # Endpoints /api/items
│       └── authController.ts      # Endpoints /api/auth
│
├── client/                       # Frontend React + Vite
│   ├── src/
│   │   ├── App.tsx              # Composant principal (page sélection + app)
│   │   ├── types.ts             # Types TypeScript frontend
│   │   ├── index.css            # Styles globaux
│   │   ├── main.tsx             # Entry point React
│   │   ├── contexts/
│   │   │   └── AuthContext.tsx  # État global authentification
│   │   └── components/
│   │       ├── AuthForm.tsx     # Formulaire login/register
│   │       └── ItemManager.tsx  # Gestion items (CRUD)
│   ├── index.html               # HTML avec meta charset UTF-8
│   ├── vite.config.ts           # Configuration Vite
│   ├── tsconfig.json            # Config TypeScript
│   └── tailwind.config.js       # Palette couleurs Tailwind
│
├── docker-compose.yml           # MongoDB container local
├── Dockerfile                   # Image Docker backend
├── package.json                 # Dépendances backend
├── tsconfig.json                # Config TypeScript backend
├── .env                         # Variables d'environnement
└── README.md                    # Documentation projet
"""

style = doc.styles['Normal']
code_style = style.font
code_style.name = 'Courier New'
code_style.size = Pt(9)

p = doc.add_paragraph(structure_text)
for run in p.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

doc.add_page_break()

# ============== 4. FICHIERS BACKEND ==============
doc.add_heading('4. Fichiers Backend', level=1)

# ---- server.ts ----
doc.add_heading('📄 server.ts - Point d\'entrée principal', level=2)
doc.add_paragraph(
    'RÔLE: Initialiser Express, configurer middleware, définir routes API, démarrer serveur\n\n'
    'Ce fichier est responsable de:\n'
    '• Créer l\'application Express\n'
    '• Configurer les middlewares (CORS, JSON parsing, UTF-8 encoding)\n'
    '• Attacher les routes API (/api/countries, /api/items, /api/auth)\n'
    '• Initialiser la base de données et seeder les pays\n'
    '• Écouter sur le port 4000\n\n'
    'CODE CLÉS:\n'
    '- Middleware UTF-8: res.setHeader(\'Content-Type\', \'application/json; charset=utf-8\')\n'
    '- Routes: app.use(\'/api/countries\', countriesRouter)\n'
    '- Startup: seedDatabase() puis app.listen(4000)'
)

# ---- country.ts ----
doc.add_heading('📄 types/country.ts - Interfaces TypeScript', level=2)
doc.add_paragraph(
    'RÔLE: Définir les types pour la sécurité du code\n\n'
    'Interfaces définies:\n'
    '• Country: code, name, flag, phoneCode\n'
    '• Link: countryCode, title, url, description\n\n'
    'EXEMPLE:\n'
    'interface Country {\n'
    '  _id?: string;\n'
    '  code: string;        // "senegal"\n'
    '  name: string;        // "Sénégal"\n'
    '  flag: string;        // "🇸🇳"\n'
    '  phoneCode: string;   // "+221"\n'
    '}'
)

# ---- subdomain.ts ----
doc.add_heading('📄 utils/subdomain.ts - Détection pays par subdomain', level=2)
doc.add_paragraph(
    'RÔLE: Extraire le code pays du host header HTTP\n\n'
    'Fonction principale: getCountryCodeFromRequest(req)\n'
    'Détecte le pattern: {country}.localhost:5173\n'
    'Utilise regex: /^([a-z]+)\\.localhost/\n\n'
    'EXEMPLE:\n'
    'Host: senegal.localhost:5173 → retourne "senegal"\n'
    'Host: afri-hub.localhost:5173 → retourne null (racine)\n\n'
    'UTILISATION:\n'
    '• Backend utilise pour isoler données par pays\n'
    '• Retourne le code dans /api/countries/current'
)

# ---- seedData.ts ----
doc.add_heading('📄 utils/seedData.ts - Initialisation base de données', level=2)
doc.add_paragraph(
    'RÔLE: Créer et remplir les collections MongoDB au démarrage\n\n'
    'Données de 10 pays africains:\n'
    '• Sénégal (🇸🇳, +221)\n'
    '• Côte d\'Ivoire (🇨🇮, +225)\n'
    '• Mali (🇲🇱, +223)\n'
    '• Ghana (🇬🇭, +233)\n'
    '• Nigeria (🇳🇬, +234)\n'
    '• Kenya (🇰🇪, +254)\n'
    '• Cameroun (🇨🇲, +237)\n'
    '• Bénin (🇧🇯, +229)\n'
    '• Congo (🇨🇬, +242)\n'
    '• Ouganda (🇺🇬, +256)\n\n'
    'FONCTION CLÉE:\n'
    'seedDatabase() - Nettoie les collections, puis insère les 10 pays'
)

# ---- countriesController.ts ----
doc.add_heading('📄 controllers/countriesController.ts - API Pays', level=2)
doc.add_paragraph(
    'RÔLE: Définir les endpoints REST pour les pays\n\n'
    'Endpoints:\n'
    '• GET /api/countries → Retourne les 10 pays\n'
    '• GET /api/countries/current → Retourne pays du subdomain détecté\n'
    '• GET /api/countries/:code → Retourne un pays spécifique\n'
    '• GET /api/countries/:code/links → Retourne liens d\'un pays\n\n'
    'SÉCURITÉ:\n'
    '• Mapping local flagMap garantit que les emojis s\'affichent\n'
    '• serializeCountry() convertit MongoDB ObjectId en hex string\n'
    '• Double fallback: backend map + frontend map'
)

# ---- authController.ts ----
doc.add_heading('📄 controllers/authController.ts - Authentification', level=2)
doc.add_paragraph(
    'RÔLE: Gérer login, register, validation JWT\n\n'
    'Endpoints:\n'
    '• POST /api/auth/login → Email + password → Retourne JWT token + user\n'
    '• POST /api/auth/register → Email + password → Crée compte + retourne token\n\n'
    'SÉCURITÉ:\n'
    '• Mots de passe hashés avec bcrypt\n'
    '• JWT signé avec secret\n'
    '• Middleware vérifie token sur /api/items'
)

# ---- itemsController.ts ----
doc.add_heading('📄 controllers/itemsController.ts - Gestion items', level=2)
doc.add_paragraph(
    'RÔLE: CRUD items (tâches) isolés par pays + utilisateur\n\n'
    'Endpoints:\n'
    '• GET /api/items → Items de l\'utilisateur actuel\n'
    '• POST /api/items → Créer nouvel item\n'
    '• PUT /api/items/:id → Modifier item\n'
    '• DELETE /api/items/:id → Supprimer item\n\n'
    'ISOLATION:\n'
    '• Items filtrés par userId + countryCode\n'
    '• countryCode extrait du subdomain\n'
    '• Chaque pays a ses propres items'
)

doc.add_page_break()

# ============== 5. FICHIERS FRONTEND ==============
doc.add_heading('5. Fichiers Frontend', level=1)

# ---- App.tsx ----
doc.add_heading('📄 client/src/App.tsx - Composant principal', level=2)
doc.add_paragraph(
    'RÔLE: Afficher page de sélection de pays (racine) OU app authentifiée (subdomain)\n\n'
    'ÉTATS GÉRÉS:\n'
    '• country: Pays détecté du subdomain\n'
    '• availableCountries: Liste des 10 pays\n'
    '• countrySearch: Texte de recherche\n'
    '• isAuthenticated: Token JWT existe?\n'
    '• activeTab: dashboard/items/profile\n\n'
    'LOGIQUE CLÉE:\n'
    'const isRootHost = hostname === \'afri-hub.localhost\'\n'
    'if (isRootHost) → Afficher page sélection\n'
    'else if (!isAuthenticated) → Afficher login\n'
    'else → Afficher dashboard\n\n'
    'FLAG_MAP: Mapping local des drapeaux pour garantir affichage'
)

# ---- AuthContext.tsx ----
doc.add_heading('📄 client/src/contexts/AuthContext.tsx - État global Auth', level=2)
doc.add_paragraph(
    'RÔLE: Gérer l\'authentification globalement (accessible partout via useAuth)\n\n'
    'FONCTIONS:\n'
    '• login(email, password): Appel /api/auth/login, stocke token\n'
    '• register(email, password): Appel /api/auth/register\n'
    '• logout(): Supprime token du localStorage\n'
    '• fetchCurrentCountry(): Appel /api/countries/current\n\n'
    'DONNÉES STOCKÉES:\n'
    '• token: JWT stocké en localStorage\n'
    '• user: {_id, email}\n'
    '• country: Pays du subdomain\n'
    '• error: Message erreur\n\n'
    'UTILISÉ PAR:\n'
    '• App.tsx: const { isAuthenticated, token, user } = useAuth()\n'
    '• ItemManager.tsx: Ajouter Authorization header au fetch'
)

# ---- AuthForm.tsx ----
doc.add_heading('📄 client/src/components/AuthForm.tsx - Formulaire Auth', level=2)
doc.add_paragraph(
    'RÔLE: UI pour login/register avec affichage drapeau + code téléphone\n\n'
    'CHAMPS:\n'
    '• Email: input type="email"\n'
    '• Password: input type="password"\n'
    '• Toggle: Bouton pour basculer Login ↔ Register\n\n'
    'AFFICHAGE:\n'
    '• Header: 🇸🇳 Sénégal +221 (pays détecté)\n'
    '• Drapeau emoji + code téléphone badge\n'
    '• Erreur: Message rouge avec icône warning\n\n'
    'STYLE:\n'
    '• Gradients emerald-amber\n'
    '• Rounded-full buttons\n'
    '• Tailwind CSS personnalisé'
)

# ---- ItemManager.tsx ----
doc.add_heading('📄 client/src/components/ItemManager.tsx - Gestion items', level=2)
doc.add_paragraph(
    'RÔLE: Afficher, créer, modifier, supprimer items (tâches)\n\n'
    'FONCTIONNALITÉS:\n'
    '• Liste des items avec checkbox completed\n'
    '• Input pour créer nouvel item\n'
    '• Bouton delete par item\n'
    '• Auto-sync avec backend\n\n'
    'API CALLS:\n'
    '• GET /api/items → Charger items utilisateur\n'
    '• POST /api/items → Créer item\n'
    '• PUT /api/items/:id → Cocher/décocher\n'
    '• DELETE /api/items/:id → Supprimer\n\n'
    'ISOLATION:\n'
    '• Items filtrés par userId + countryCode\n'
    '• Chaque utilisateur ne voit que ses items\n'
    '• Chaque pays a ses propres items'
)

# ---- types.ts ----
doc.add_heading('📄 client/src/types.ts - Types Frontend', level=2)
doc.add_paragraph(
    'RÔLE: Définir interfaces TypeScript pour le frontend\n\n'
    'INTERFACES:\n'
    '• Country: {code, name, flag, phoneCode}\n'
    '• User: {_id, email}\n'
    '• Item: {_id, title, completed, userId, countryCode}\n'
    '• AuthResponse: {token, user}\n\n'
    'IMPORTÉ PAR:\n'
    '• App.tsx\n'
    '• AuthContext.tsx\n'
    '• ItemManager.tsx'
)

doc.add_page_break()

# ============== 6. FICHIERS CONFIGURATION ==============
doc.add_heading('6. Fichiers de configuration', level=1)

# ---- docker-compose.yml ----
doc.add_heading('📄 docker-compose.yml - MongoDB local', level=2)
doc.add_paragraph(
    'RÔLE: Démarrer MongoDB en container Docker\n\n'
    'CONFIGURATION:\n'
    '• Image: mongo:7.0\n'
    '• Port: 27017 (standard MongoDB)\n'
    '• Root user: admin / password\n'
    '• Database: afrihub\n'
    '• Volumes: mongo_data pour persister\n\n'
    'DÉMARRAGE:\n'
    'docker-compose up -d\n\n'
    'ARRÊT:\n'
    'docker-compose down'
)

# ---- .env ----
doc.add_heading('📄 .env - Variables d\'environnement', level=2)
doc.add_paragraph(
    'RÔLE: Configurer les paramètres du projet\n\n'
    'VARIABLES:\n'
    '• MONGO_URI: mongodb://admin:password@localhost:27017/afrihub?authSource=admin\n'
    '• MONGODB_DB: afrihub\n'
    '• PORT: 4000\n'
    '• NODE_ENV: development\n'
    '• JWT_SECRET: Clé pour signer tokens\n\n'
    'IMPORTANT:\n'
    '• Ne jamais committer .env en production\n'
    '• Adapter MONGO_URI pour production (MongoDB Atlas)'
)

# ---- package.json ----
doc.add_heading('📄 package.json - Dépendances', level=2)
doc.add_paragraph(
    'RÔLE: Déclarer dépendances npm et scripts\n\n'
    'SCRIPTS:\n'
    '• npm run dev: Démarre avec nodemon (hot reload)\n'
    '• npm run build: Compile TypeScript\n'
    '• npm start: Démarre production\n\n'
    'DÉPENDANCES:\n'
    '• express: Web framework\n'
    '• mongodb: Driver MongoDB\n'
    '• bcrypt: Hash mots de passe\n'
    '• jsonwebtoken: JWT tokens\n'
    '• cors: Cross-origin requests\n'
    '• dotenv: Variables d\'environnement'
)

# ---- tsconfig.json ----
doc.add_heading('📄 tsconfig.json - Configuration TypeScript', level=2)
doc.add_paragraph(
    'RÔLE: Configurer le compilateur TypeScript\n\n'
    'OPTIONS CLÉES:\n'
    '• target: ES2020\n'
    '• module: ESNext\n'
    '• strict: true (type checking strict)\n'
    '• moduleResolution: bundler\n'
    '• skipLibCheck: true (ignore erreurs lib)\n\n'
    'RÉSULTAT:\n'
    '• Code type-safe\n'
    '• Erreurs détectées à la compilation\n'
    '• Auto-completion IDE complète'
)

# ---- vite.config.ts ----
doc.add_heading('📄 client/vite.config.ts - Configuration Vite', level=2)
doc.add_paragraph(
    'RÔLE: Configurer le bundler Vite\n\n'
    'CONFIGURATION:\n'
    '• Port: 5173\n'
    '• React plugin: @vitejs/plugin-react-swc\n'
    '• SWC: Compilateur Rust rapide\n'
    '• HMR: Hot module replacement\n\n'
    'RÉSULTAT:\n'
    '• Serveur dev super rapide\n'
    '• Build production optimisé\n'
    '• Import ES modules nativement'
)

# ---- tailwind.config.js ----
doc.add_heading('📄 client/tailwind.config.js - Design System', level=2)
doc.add_paragraph(
    'RÔLE: Configurer Tailwind CSS et palette de couleurs\n\n'
    'COULEURS PRINCIPALES:\n'
    '• Emerald-600 (#059669): Boutons, badges\n'
    '• Amber-500 (#F59E0B): Gradient logo\n'
    '• Slate-900 (#0F172A): Texte principal\n'
    '• Slate-600 (#475569): Texte secondaire\n'
    '• Slate-200 (#E2E8F0): Bordures\n\n'
    'RÉSULTAT:\n'
    '• Palette cohérente\n'
    '• Classes utilitaires (px-4, py-3, etc.)\n'
    '• Responsive design (sm:, md:, lg:)'
)

doc.add_page_break()

# ============== 7. EXPLICATIONS DÉTAILLÉES ==============
doc.add_heading('7. Explications détaillées du code', level=1)

# Flux recherche
doc.add_heading('🔍 Flux: Recherche d\'un pays', level=2)
doc.add_paragraph(
    '1. Utilisateur tape "sén" dans l\'input\n'
    '   → onChange déclenche setCountrySearch("sén")\n\n'
    '2. useMemo recalcule filteredCountries:\n'
    '   countries.filter(c => \n'
    '     c.name.toLowerCase().includes("sén") OR\n'
    '     c.phoneCode.includes("sén") OR\n'
    '     c.code.includes("sén")\n'
    '   )\n\n'
    '3. Résultats affichés si countrySearch !== ""\n'
    '   → "Sénégal" trouvé\n\n'
    '4. Au clic:\n'
    '   handleSelectCountry({code: "senegal", ...})\n'
    '   → window.location.href = "senegal.localhost:5173"\n\n'
    '5. Redirection vers subdomain'
)

# Flux authentification
doc.add_heading('🔐 Flux: Authentification', level=2)
doc.add_paragraph(
    '1. Utilisateur arrive senegal.localhost:5173 (après sélection)\n\n'
    '2. App.tsx détecte isRootHost === false\n'
    '   → Lance loadCountry()\n\n'
    '3. fetch /api/countries/current\n'
    '   Backend: getCountryCodeFromRequest(req)\n'
    '   → Regex extrait "senegal" du host header\n'
    '   → Retourne pays Sénégal avec drapeau 🇸🇳 + +221\n\n'
    '4. setCountry(sénégal)\n'
    '   → Affiche drapeau + nom + téléphone\n\n'
    '5. Affiche AuthForm (login/register)\n\n'
    '6. Utilisateur rentre email + password\n'
    '   → fetch /api/auth/login\n'
    '   → Backend crée JWT token\n'
    '   → localStorage.setItem("authToken", token)\n\n'
    '7. Affiche dashboard (items, profile, etc.)'
)

# Flux items isolés
doc.add_heading('📝 Flux: Items isolés par pays', level=2)
doc.add_paragraph(
    '1. Utilisateur connecté sur senegal.localhost:5173\n'
    '   → AuthContext.token existe\n'
    '   → countryCode = "senegal" (extrait du subdomain)\n\n'
    '2. ItemManager.tsx chargera les items:\n'
    '   fetch /api/items\n'
    '   Headers: Authorization: Bearer {token}\n\n'
    '3. Backend filtre:\n'
    '   items.find({\n'
    '     userId: {id de l\'utilisateur},\n'
    '     countryCode: "senegal"\n'
    '   })\n\n'
    '4. Affiche uniquement les items de cet utilisateur EN SENEGAL\n\n'
    '5. Si utilisateur va sur kenya.localhost:5173:\n'
    '   → countryCode change à "kenya"\n'
    '   → Items différents chargés\n'
    '   → Isolation complète par pays'
)

doc.add_page_break()

# ============== RÉSUMÉ ==============
doc.add_heading('📊 Résumé du projet', level=1)

# Table résumé
summary_table = doc.add_table(rows=1, cols=2)
summary_table.style = 'Light Grid Accent 1'
hdr_cells = summary_table.rows[0].cells
hdr_cells[0].text = 'Aspect'
hdr_cells[1].text = 'Details'

summary_data = [
    ('Pays supportés', '10 (Sénégal, Côte d\'Ivoire, Mali, Ghana, Nigeria, Kenya, Cameroun, Bénin, Congo, Ouganda)'),
    ('Endpoints API', '12+'),
    ('Types TypeScript', '8+'),
    ('Composants React', '5+'),
    ('Fonctionnalités', '7 (sélection pays, search, multi-tenant, auth, items, profil)'),
    ('Port frontend', '5173'),
    ('Port backend', '4000'),
    ('Port MongoDB', '27017'),
    ('Palette couleurs', 'Emerald/Amber/Slate (Tailwind)'),
    ('Drapeaux', 'Emoji natifs (🇸🇳, 🇨🇮, etc.)'),
]

for aspect, detail in summary_data:
    row_cells = summary_table.add_row().cells
    row_cells[0].text = aspect
    row_cells[1].text = detail

# ============== NOTES ==============
doc.add_page_break()
doc.add_heading('📝 Notes importantes', level=1)

notes = [
    'Type-Safe: Tout en TypeScript avec strict mode',
    'Multi-tenant: Isolation complète par subdomain',
    'Scalable: Code prêt pour 50+ pays',
    'Moderne: React 18 + Vite + Tailwind CSS 3',
    'Sécurisé: JWT tokens + bcrypt hash',
    'Responsive: Mobile-first design',
    'Production-ready: Docker + environment variables'
]

for note in notes:
    p = doc.add_paragraph(note, style='List Bullet')

# ============== SAUVEGARDER ==============
doc.save('c:\\Users\\USER\\projects\\AFRI-HUB\\AFRI_HUB_DOCUMENTATION.docx')
print('✅ Fichier Word créé: AFRI_HUB_DOCUMENTATION.docx')
print('📂 Localisation: c:\\Users\\USER\\projects\\AFRI-HUB\\')
